"""
管理画面操作マニュアル DOCX 生成スクリプト

Usage:
    python3 docs/manual/generate_docx.py

前提:
    pip install python-docx
    node docs/manual/capture.mjs  (スクリーンショット撮影済み)
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

SCRIPT_DIR = Path(__file__).parent
SCREENSHOT_DIR = SCRIPT_DIR / "screenshots"
OUTPUT_PATH = SCRIPT_DIR / "管理画面操作マニュアル.docx"


def set_cell_shading(cell, color_hex):
    """セルの背景色を設定する"""
    shading = cell._element.get_or_add_tcPr()
    shading_elem = shading.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): color_hex,
        },
    )
    shading.append(shading_elem)


def add_table(doc, headers, rows, col_widths=None):
    """テーブルを追加する"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ヘッダー行
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        set_cell_shading(cell, "F5F5F5")

    # データ行
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    # 列幅
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(width)

    doc.add_paragraph("")
    return table


def add_image(doc, subdir, filename, width=Inches(5.5)):
    """スクリーンショットを追加する"""
    img_path = SCREENSHOT_DIR / subdir / filename
    if img_path.exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(img_path), width=width)
    else:
        doc.add_paragraph(f"[画像: {filename} が見つかりません]")


def add_step(doc, step_num, title, description):
    """手順を追加する"""
    p = doc.add_paragraph()
    run = p.add_run(f"手順 {step_num}: {title}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if description:
        doc.add_paragraph(description)


def build_document():
    """DOCX ドキュメントを生成する"""
    doc = Document()

    # ページ設定（A4）
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # デフォルトフォント
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Yu Gothic"
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")

    # ========================================
    # 表紙
    # ========================================
    for _ in range(6):
        doc.add_paragraph("")

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("管理画面 操作マニュアル")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title_run.font.color.rgb = RGBColor(0xE6, 0x8A, 0x00)

    doc.add_paragraph("")

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_p.add_run("ちほプロジェクト - みんなの平泉")
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph("")
    doc.add_paragraph("")

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_p.add_run(
        "対象: 管理者（admin）/ 役場職員（municipal）\n更新日: 2026-03-07"
    )
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_page_break()

    # ========================================
    # 目次
    # ========================================
    doc.add_heading("目次", level=1)

    toc_items = [
        "第1章 管理者パネル（/admin）",
        "  1. ログイン",
        "  2. ダッシュボード",
        "  3. ユーザー管理",
        "  4. 応募管理",
        "  5. 操作ログ",
        "",
        "第2章 役場パネル（/municipal）",
        "  6. ログイン",
        "  7. 応募一覧（閲覧）",
        "  8. CSVダウンロード",
        "  9. 操作ログ（閲覧）",
        "",
        "付録A データ項目仕様",
        "  A-1. 管理者パネル一覧画面の表示項目",
        "  A-2. CSVファイル出力項目",
    ]
    for item in toc_items:
        if item == "":
            doc.add_paragraph("")
        elif item.startswith("第") or item.startswith("付録"):
            p = doc.add_paragraph()
            run = p.add_run(item)
            run.bold = True
            run.font.size = Pt(11)
        else:
            doc.add_paragraph(item, style="List Bullet")

    doc.add_page_break()

    # ========================================
    # 第1章 管理者パネル
    # ========================================
    doc.add_heading("第1章 管理者パネル（/admin）", level=1)
    doc.add_paragraph(
        "管理者は、ユーザー管理・応募管理・操作ログの閲覧など、システム全体を管理できます。"
    )

    # --- 1. ログイン ---
    doc.add_heading("1. 管理者パネルへのログイン", level=2)
    doc.add_paragraph("URL: https://サイトURL/admin/login")

    add_step(doc, "1-1", "ログイン画面を開く", "ブラウザで管理パネルのURLにアクセスすると、ログイン画面が表示されます。")
    add_image(doc, "admin", "01_login_empty.png")

    add_step(doc, "1-2", "メールアドレスを入力する", "赤枠のメールアドレス欄に、管理者アカウントのメールアドレスを入力します。")
    add_image(doc, "admin", "02_login_email_highlight.png")

    add_step(doc, "1-3", "パスワードを入力し、ログインボタンを押す", "パスワードを入力し、赤枠の「ログイン」ボタンをクリックします。")
    add_image(doc, "admin", "03_login_filled.png")

    # --- 2. ダッシュボード ---
    doc.add_heading("2. ダッシュボード", level=2)
    doc.add_paragraph("ログインすると、ダッシュボードが表示されます。左側のサイドバーから各機能にアクセスできます。")
    add_image(doc, "admin", "04_dashboard.png")

    doc.add_paragraph("サイドバーのメニュー:")
    add_table(
        doc,
        ["メニュー", "機能"],
        [
            ["ダッシュボード", "トップページ"],
            ["応募管理", "全応募の一覧・詳細・編集・CSVダウンロード"],
            ["Filament Shield > ロール", "ロール・権限の管理"],
            ["システム管理 > ユーザー管理", "全ユーザーの一覧・作成・編集・削除"],
            ["システム管理 > 操作ログ", "操作履歴の閲覧"],
        ],
        col_widths=[5, 11],
    )

    # --- 3. ユーザー管理 ---
    doc.add_heading("3. ユーザー管理", level=2)

    add_step(doc, "3-1", "ユーザー管理を開く", "サイドバーの赤枠「ユーザー管理」をクリックします。")
    add_image(doc, "admin", "05_sidebar_users_highlight.png")

    add_step(doc, "3-2", "ユーザー一覧", "全ユーザーの一覧が表示されます。氏名・メールアドレス・ロールで検索・フィルタリングが可能です。")
    add_image(doc, "admin", "06_users_list.png")

    doc.add_paragraph("操作:")
    ops = [
        "「表示」: ユーザーの詳細を閲覧",
        "「編集」: ユーザー情報を変更",
        "検索欄: 氏名・メールアドレスで検索",
        "フィルタアイコン: ロールでフィルタリング",
    ]
    for op in ops:
        doc.add_paragraph(op, style="List Bullet")

    add_step(doc, "3-3", "ユーザーを新規作成する", "赤枠の「作成」ボタンをクリックします。")
    add_image(doc, "admin", "07_users_new_button_highlight.png")

    doc.add_paragraph("ユーザー作成フォームが表示されます。必要な情報を入力して「作成」ボタンを押します。")
    add_image(doc, "admin", "08_users_create_form.png")

    doc.add_paragraph("入力項目:")
    add_table(
        doc,
        ["項目", "必須", "説明"],
        [
            ["氏名", "○", "ユーザーの名前"],
            ["メールアドレス", "○", "ログイン用（重複不可）"],
            ["パスワード", "○", "8文字以上"],
            ["ロール", "○", "ワーカー / 企業 / 管理者 / 役所 から選択"],
            ["Spatieロール", "-", "権限管理用ロール（複数選択可）"],
        ],
        col_widths=[4, 2, 10],
    )

    add_step(doc, "3-4", "ユーザー詳細を表示する", "一覧画面で赤枠の行をクリック、または「表示」リンクをクリックします。")
    add_image(doc, "admin", "09_users_row_highlight.png")
    doc.add_paragraph("ユーザーの詳細情報が表示されます。")
    add_image(doc, "admin", "10_users_view.png")

    # --- 4. 応募管理 ---
    doc.add_heading("4. 応募管理", level=2)

    add_step(doc, "4-1", "応募管理を開く", "サイドバーの赤枠「応募管理」をクリックします。")
    add_image(doc, "admin", "13_sidebar_applications_highlight.png")

    add_step(doc, "4-2", "応募一覧", "全応募データの一覧が表示されます。")
    add_image(doc, "admin", "14_applications_list.png")

    add_step(doc, "4-3", "CSVダウンロード", "赤枠の「CSVダウンロード」ボタンをクリックすると、応募データをCSVファイルとしてダウンロードできます。")
    add_image(doc, "admin", "15_applications_csv_highlight.png")

    add_step(doc, "4-4", "応募詳細を表示する", "一覧の赤枠の行をクリック、または「表示」リンクをクリックします。")
    add_image(doc, "admin", "16_applications_row_highlight.png")
    doc.add_paragraph("応募の詳細情報が表示されます。")
    add_image(doc, "admin", "17_applications_view.png")

    # --- 5. 操作ログ ---
    doc.add_heading("5. 操作ログ", level=2)
    doc.add_paragraph("システムの操作履歴を閲覧できます。ユーザーの登録・編集・応募のステータス変更などが記録されています。")

    add_step(doc, "5-1", "操作ログを開く", "サイドバーの赤枠「操作ログ」をクリックします。")
    add_image(doc, "admin", "20_sidebar_logs_highlight.png")

    add_step(doc, "5-2", "操作ログ一覧", "操作ログの一覧が表示されます。")
    add_image(doc, "admin", "21_activity_logs_list.png")

    add_step(doc, "5-3", "操作ログ詳細", "一覧の赤枠の行をクリックすると、詳細が表示されます。")
    add_image(doc, "admin", "22_activity_logs_row_highlight.png")
    doc.add_paragraph("操作の詳細（変更内容・変更前後の値など）が確認できます。")
    add_image(doc, "admin", "23_activity_logs_view.png")

    doc.add_page_break()

    # ========================================
    # 第2章 役場パネル
    # ========================================
    doc.add_heading("第2章 役場パネル（/municipal）", level=1)
    doc.add_paragraph("役場職員は、応募データの閲覧とCSVダウンロードのみ利用できます。データの編集・削除はできません。")

    # --- 6. ログイン ---
    doc.add_heading("6. 役場パネルへのログイン", level=2)
    doc.add_paragraph("URL: https://サイトURL/municipal/login")

    add_step(doc, "6-1", "ログイン画面を開く", "ブラウザで役場パネルのURLにアクセスすると、ログイン画面が表示されます。")
    add_image(doc, "municipal", "01_login_empty.png")

    add_step(doc, "6-2", "メールアドレスを入力する", "赤枠のメールアドレス欄に、役場アカウントのメールアドレスを入力します。")
    add_image(doc, "municipal", "02_login_email_highlight.png")

    add_step(doc, "6-3", "パスワードを入力し、ログインボタンを押す", "パスワードを入力し、赤枠の「ログイン」ボタンをクリックします。")
    add_image(doc, "municipal", "03_login_filled.png")
    doc.add_paragraph("ログイン後、ダッシュボードが表示されます。")
    add_image(doc, "municipal", "04_dashboard.png")

    # --- 7. 応募一覧 ---
    doc.add_heading("7. 応募一覧（閲覧）", level=2)

    add_step(doc, "7-1", "応募一覧を開く", "サイドバーの「応募管理」をクリックします。")

    add_step(doc, "7-2", "応募一覧（読取専用）", "全応募データの一覧が表示されます。閲覧のみで、編集・削除はできません。")
    add_image(doc, "municipal", "06_applications_list.png")

    add_step(doc, "7-3", "応募詳細を表示する", "赤枠の行をクリックすると、応募の詳細が表示されます。")
    add_image(doc, "municipal", "07_applications_row_highlight.png")
    doc.add_paragraph("応募の詳細情報（応募者・募集内容・ステータスなど）が確認できます。")
    add_image(doc, "municipal", "08_applications_view.png")

    # --- 8. CSVダウンロード ---
    doc.add_heading("8. CSVダウンロード", level=2)
    doc.add_paragraph("応募データをCSVファイルとしてダウンロードできます。期間を指定してフィルタリングが可能です。")

    add_step(doc, "8-1", "CSVダウンロード画面を開く", "サイドバーの赤枠「CSVダウンロード」をクリックします。")
    add_image(doc, "municipal", "09_sidebar_csv_highlight.png")

    add_step(doc, "8-2", "期間を指定する（任意）", "開始日・終了日を入力して、ダウンロード対象の期間を指定できます。未指定の場合は全期間が対象です。")
    add_image(doc, "municipal", "10_csv_download_page.png")

    add_step(doc, "8-3", "ダウンロードを実行する", "赤枠の「CSVダウンロード」ボタンをクリックすると、CSVファイルがダウンロードされます。")
    add_image(doc, "municipal", "11_csv_download_button_highlight.png")

    # --- 9. 操作ログ ---
    doc.add_heading("9. 操作ログ（閲覧）", level=2)
    doc.add_paragraph("システムの操作履歴を閲覧できます。読取専用です。")

    add_step(doc, "9-1", "操作ログを開く", "サイドバーの赤枠「操作ログ」をクリックします。")
    add_image(doc, "municipal", "12_sidebar_logs_highlight.png")

    add_step(doc, "9-2", "操作ログ一覧", "操作ログの一覧が表示されます。")
    add_image(doc, "municipal", "13_activity_logs_list.png")

    add_step(doc, "9-3", "操作ログ詳細", "赤枠の行をクリックすると、操作の詳細が表示されます。")
    add_image(doc, "municipal", "14_activity_logs_row_highlight.png")
    doc.add_paragraph("操作の詳細（変更内容・変更前後の値など）が確認できます。")
    add_image(doc, "municipal", "15_activity_logs_view.png")

    doc.add_page_break()

    # ========================================
    # 付録A データ項目仕様
    # ========================================
    doc.add_heading("付録A データ項目仕様", level=1)

    # --- A-1. 管理者パネル一覧画面の表示項目 ---
    doc.add_heading("A-1. 管理者パネル一覧画面の表示項目", level=2)

    # ユーザー管理一覧
    doc.add_heading("A-1-1. ユーザー管理一覧（/admin/users）", level=3)
    add_table(
        doc,
        ["表示カラム名", "テーブル名", "フィールド名", "サンプル"],
        [
            ["氏名", "users", "name", "森田 強"],
            ["メールアドレス", "users", "email", "t.morita@example.com"],
            ["ロール", "users", "role", "ワーカー（worker）"],
            ["Spatieロール", "model_has_roles → roles", "roles.name", "worker"],
            ["登録日", "users", "created_at", "2026/03/05"],
        ],
        col_widths=[3.5, 4, 4, 4.5],
    )

    # 応募管理一覧
    doc.add_heading("A-1-2. 応募管理一覧（/admin/job-applications）", level=3)
    add_table(
        doc,
        ["表示カラム名", "テーブル名", "フィールド名", "サンプル"],
        [
            ["求人タイトル", "job_posts", "job_title", "農作業のお手伝い"],
            ["企業名", "users（company）", "name", "平泉観光協会"],
            ["応募者名", "users（worker）", "name", "森田 強"],
            ["メールアドレス", "users（worker）", "email", "worker@example.com"],
            ["ステータス", "job_applications", "status", "応募中（applied）"],
            ["応募日", "job_applications", "applied_at", "2026/03/05"],
            ["判定日", "job_applications", "judged_at", "2026/03/06"],
        ],
        col_widths=[3.5, 4, 4, 4.5],
    )
    p = doc.add_paragraph()
    run = p.add_run("ステータスの値:")
    run.bold = True
    add_table(
        doc,
        ["DB値", "画面表示", "バッジ色", "説明"],
        [
            ["applied", "応募中", "黄色", "応募直後の状態"],
            ["accepted", "承認", "緑色", "ホストが応募を承認した状態"],
            ["rejected", "不承認", "赤色", "ホストが応募を不承認とした状態"],
        ],
        col_widths=[3, 3, 3, 7],
    )

    # 操作ログ一覧
    doc.add_heading("A-1-3. 操作ログ一覧（/admin/activity-logs）", level=3)
    add_table(
        doc,
        ["表示カラム名", "テーブル名", "フィールド名", "サンプル"],
        [
            ["日時", "activity_log", "created_at", "2026/03/05 14:30:00"],
            ["操作者", "activity_log → users", "causer.name", "Admin"],
            ["対象モデル", "activity_log", "subject_type", "User"],
            ["操作内容", "activity_log", "event", "作成（created）"],
            ["対象ID", "activity_log", "subject_id", "15"],
        ],
        col_widths=[3.5, 4, 4, 4.5],
    )
    p = doc.add_paragraph()
    run = p.add_run("操作内容の値:")
    run.bold = True
    add_table(
        doc,
        ["DB値", "画面表示", "バッジ色"],
        [
            ["created", "作成", "緑色"],
            ["updated", "更新", "黄色"],
            ["deleted", "削除", "赤色"],
        ],
        col_widths=[4, 4, 4],
    )

    doc.add_page_break()

    # --- A-2. CSVファイル出力項目 ---
    doc.add_heading("A-2. CSVファイル出力項目", level=2)
    doc.add_paragraph(
        "管理者パネルの応募管理一覧、および役場パネルのCSVダウンロード画面からダウンロードできるCSVファイルの項目です。"
    )
    doc.add_paragraph(
        "ファイル名の形式: 応募一覧_YYYYMMDD_HHmmss.csv（例: 応募一覧_20260307_143000.csv）"
    )
    doc.add_paragraph("文字コード: UTF-8（BOM付き、Excel対応）")
    doc.add_paragraph("")

    add_table(
        doc,
        ["#", "CSVカラム名", "テーブル名", "フィールド名", "サンプル"],
        [
            ["1", "氏名", "users", "name", "森田 強"],
            ["2", "メールアドレス", "users", "email", "worker@example.com"],
            ["3", "ハンドルネーム", "worker_profiles", "handle_name", "もりたん"],
            ["4", "性別", "worker_profiles", "gender", "男性（male）"],
            ["5", "生年月日", "worker_profiles", "birthdate", "1990/05/15"],
            ["6", "現住所", "worker_profiles", "current_address", "平泉町平泉字衣関1-2"],
            ["7", "電話番号", "worker_profiles", "phone_number", "090-1234-5678"],
            ["8", "出身地", "locations（birth_location）", "prefecture", "岩手県"],
            ["9", "居住地（都道府県）", "locations（current_location_1）", "prefecture", "岩手県"],
            ["10", "志望動機", "job_applications", "motive", "地域活動に興味があり..."],
            ["11", "応募理由", "job_applications", "reasons（JSON）", "力仕事ができる、写真撮影"],
            ["12", "ステータス", "job_applications", "status", "応募中"],
            ["13", "応募日", "job_applications", "applied_at", "2026/03/05 14:30"],
            ["14", "判定日", "job_applications", "judged_at", "2026/03/06 10:00"],
            ["15", "求人タイトル", "job_posts", "job_title", "農作業のお手伝い"],
            ["16", "企業名", "users（company）", "name", "平泉観光協会"],
        ],
        col_widths=[1, 3.5, 3.5, 4, 4],
    )

    p = doc.add_paragraph()
    run = p.add_run("CSV出力時の変換ルール:")
    run.bold = True
    add_table(
        doc,
        ["フィールド", "DB値", "CSV出力値"],
        [
            ["性別", "male", "男性"],
            ["性別", "female", "女性"],
            ["性別", "other", "その他"],
            ["ステータス", "applied", "応募中"],
            ["ステータス", "accepted", "承認"],
            ["ステータス", "rejected", "不承認"],
            ["応募理由", 'JSON配列 ["理由1","理由2"]', "理由1、理由2（読点区切り）"],
            ["日時", "2026-03-05 14:30:00", "2026/03/05 14:30"],
            ["生年月日", "1990-05-15", "1990/05/15"],
        ],
        col_widths=[3.5, 5, 5],
    )

    doc.add_page_break()

    # ========================================
    # 補足
    # ========================================
    doc.add_heading("補足", level=1)

    doc.add_heading("権限の違い", level=2)
    add_table(
        doc,
        ["操作", "管理者", "役場"],
        [
            ["ユーザーの作成・編集・削除", "○", "-"],
            ["応募の閲覧", "○", "○"],
            ["応募のステータス変更", "○", "-"],
            ["CSVダウンロード", "○", "○"],
            ["操作ログの閲覧", "○", "○"],
            ["ロール・権限の管理", "○", "-"],
        ],
        col_widths=[8, 4, 4],
    )

    doc.add_heading("スクリーンショットの再撮影", level=2)
    doc.add_paragraph("画面の変更があった場合は、以下のコマンドでスクリーンショットを再撮影し、docxを再生成できます。")
    p = doc.add_paragraph()
    run = p.add_run("node docs/manual/capture.mjs\npython3 docs/manual/generate_docx.py")
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    # 保存
    doc.save(str(OUTPUT_PATH))
    print(f"生成完了: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
